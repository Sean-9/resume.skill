#!/usr/bin/env python3
"""JSON → docx。

用法：
    python render.py <简历.json> <输出.docx> [--template assets/template.docx]

依赖：python-docx。模板可选；缺省时新建空白文档。
排版规则见 assets/typography.md —— 本脚本只做排版，不生成任何简历内容。
中英同源字体，三级字号、三档间距、日期右对齐、论文标题 note 降级，全部按 typography.md 实现。
"""

import argparse
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# --- typography.md 数值 ---
FONT_NAME = "Noto Sans CJK SC"
SIZE_NAME = 20
SIZE_POSITIONING = 10.5
SIZE_CONTACT = 9
SIZE_SECTION = 14
SIZE_ENTRY = 11.5
SIZE_BULLET = 10
SIZE_NOTE = 9
SPACE_SECTION = Pt(6.5)
SPACE_ENTRY = Pt(2.5)
SPACE_ENTRY_HEADER = Pt(2.0)
SPACE_BULLET = Pt(1.0)
SPACE_NOTE = Pt(1.5)
GRAY = RGBColor(0x59, 0x59, 0x59)
DARK = RGBColor(0x21, 0x21, 0x21)


def _apply_font(run, size=None, bold=None, color=None):
    # 中英同源：ascii 与 eastAsia 同时设为同一字体
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _para(doc, text, size, bold=False, align=None, color=DARK,
          space_before=None, space_after=None, tab_stops=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    for pos in (tab_stops or []):
        p.paragraph_format.tab_stops.add_tab_stop(pos[0], pos[1])
    r = p.add_run(text)
    _apply_font(r, size=size, bold=bold, color=color)
    return p


def _section_title(doc, text, tab_stops):
    p = _para(doc, text, SIZE_SECTION, bold=True, color=DARK,
              space_before=SPACE_SECTION, space_after=Pt(2.5), tab_stops=tab_stops)
    # 标题下缘细线
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _entry_header(doc, title, subtitle, dates, right_tab, space_before):
    # 标题 + 角色靠左，日期用右对齐制表位钉到右边界
    text = title
    if subtitle:
        text += "　" + subtitle
    p = _para(doc, "", SIZE_ENTRY, bold=True, space_before=space_before,
              space_after=SPACE_ENTRY_HEADER, tab_stops=[(right_tab, WD_TAB_ALIGNMENT.RIGHT)])
    r1 = p.add_run(text)
    _apply_font(r1, size=SIZE_ENTRY, bold=True, color=DARK)
    if dates:
        r2 = p.add_run("\t" + dates)
        _apply_font(r2, size=SIZE_ENTRY, bold=False, color=GRAY)
    return p


def _set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(SIZE_BULLET)


def _right_tab(doc):
    # A4 页面：右对齐制表位放在内容区右边界
    sec = doc.sections[0]
    width = sec.page_width - sec.left_margin - sec.right_margin
    return sec.left_margin + width


def build(doc, data):
    _set_normal_style(doc)
    right_tab = _right_tab(doc)

    # H1 姓名
    _para(doc, data.get("name", ""), SIZE_NAME, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    # 定位语
    if data.get("positioning"):
        _para(doc, data["positioning"], SIZE_POSITIONING,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK, space_after=Pt(4))

    # 联系方式
    contact = data.get("contact", {})
    contact_parts = []
    for key in ("phone", "email", "wechat", "city", "github", "portfolio"):
        val = str(contact.get(key, "")).strip()
        if val:
            contact_parts.append(val)
    if contact_parts:
        _para(doc, "  |  ".join(contact_parts), SIZE_CONTACT,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=Pt(2))

    for i, section in enumerate(data.get("sections", [])):
        _section_title(doc, section.get("title", ""), tab_stops=[])
        if i == 0:
            # 首个章节标题与页首紧凑
            doc.paragraphs[-1].paragraph_format.space_before = Pt(4)

        for item in section.get("items", []):
            if isinstance(item, str):
                _para(doc, item, SIZE_BULLET, space_after=SPACE_ENTRY)
            else:
                _entry_header(doc, item.get("title", ""),
                              item.get("subtitle", ""), item.get("dates", ""),
                              right_tab, space_before=SPACE_ENTRY)
                if item.get("note"):
                    _para(doc, item["note"], SIZE_NOTE, color=GRAY,
                          space_after=SPACE_NOTE)
                for bullet in item.get("bullets", []):
                    bp = doc.add_paragraph(bullet, style="List Bullet")
                    bp.paragraph_format.space_after = SPACE_BULLET
                    for run in bp.runs:
                        _apply_font(run, size=SIZE_BULLET, color=DARK)


def main():
    parser = argparse.ArgumentParser(description="JSON → docx 渲染（按 typography.md）")
    parser.add_argument("input", help="简历 JSON 文件")
    parser.add_argument("output", help="输出 docx 路径")
    parser.add_argument("--template", default=None,
                        help="可选模板 docx；缺省时新建空白文档")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(args.template) if args.template and os.path.exists(args.template) else Document()
    build(doc, data)
    doc.core_properties.title = data.get("name", "简历")

    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
    doc.save(args.output)
    print(f"已生成 {args.output}")


if __name__ == "__main__":
    main()
